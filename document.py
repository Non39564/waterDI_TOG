from docx import *
import docx

def document_moc_or_auto(file):
    document = docx.Document(file)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if 'ผู้ขอใช้' in paragraph.text:
                        paragraph.text = paragraph.text.replace("ผู้ขอใช้", "{{name}}")
                    if 'เบอร์ติดต่อผู้ใช้' in paragraph.text:
                        paragraph.text = paragraph.text.replace("เบอร์ติดต่อผู้ใช้", "{{phone}}")
                    if 'ตำแหน่งผู้ใช้' in paragraph.text:
                        paragraph.text = paragraph.text.replace("ตำแหน่งผู้ใช้", "{{position}}")
                    if 'อีเมลผู้ใช้' in paragraph.text:
                        paragraph.text = paragraph.text.replace("อีเมลผู้ใช้", "{{email}}")
                    if 'แผนกผู้ใช้' in paragraph.text:
                        paragraph.text = paragraph.text.replace("แผนกผู้ใช้", "{{department}}")
                    if 'ฝ่ายผู้ใช้' in paragraph.text:
                        paragraph.text = paragraph.text.replace("ฝ่ายผู้ใช้", "{{part}}")
                    if 'วันที่ผู้ใช้' in paragraph.text:
                        paragraph.text = paragraph.text.replace("วันที่ผู้ใช้", "{{date}}")
                    if 'เวลาผู้ใช้' in paragraph.text:
                        paragraph.text = paragraph.text.replace("เวลาผู้ใช้", "{{time}}")
                    if 'สถานที่ผู้ใช้' in paragraph.text:
                        paragraph.text = paragraph.text.replace("สถานที่ผู้ใช้", "{{place}}")
                    if 'ระบบอะไร' in paragraph.text:
                        paragraph.text = paragraph.text.replace("ระบบอะไร", "{{service}}")
                    if 'เหตุผลขอใช้ระบบ' in paragraph.text:
                        paragraph.text = paragraph.text.replace("เหตุผลขอใช้ระบบ", "{{reason}}")
                    if 'ระบบปัจจุบัน' in paragraph.text:
                        paragraph.text = paragraph.text.replace("ระบบปัจจุบัน", "{{system_now}}")
                    if 'รายละเอียดผู้ใช้' in paragraph.text:
                        paragraph.text = paragraph.text.replace("รายละเอียดผู้ใช้", "{{detail}}")
                    if 'ระยะเวลาที่ต้องการให้เสร็จ ' in paragraph.text:
                        paragraph.text = paragraph.text.replace("ระยะเวลาที่ต้องการให้เสร็จ ", "{{date_end}}")
                    if 'หมายเหตุผู้ใช้' in paragraph.text:
                        paragraph.text = paragraph.text.replace("หมายเหตุผู้ใช้", "{{note}}")
                        document.save('./static/document/test.docx')


def document_it(file):
    document = docx.Document(file)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if 'ดำเนินการได้หรือไม่' in paragraph.text:
                        paragraph.text = paragraph.text.replace("ดำเนินการได้หรือไม่", "{{do}}")
                    if 'มีอะไรเกิดขึ้นหรือไม่' in paragraph.text:
                        paragraph.text = paragraph.text.replace("มีอะไรเกิดขึ้นหรือไม่", "{{detail}}")
                    if 'เสร็จสิ้นไอที' in paragraph.text:
                        paragraph.text = paragraph.text.replace("เสร็จสิ้นไอที", "{{success}}")
                    if 'วันที่เสร็จสิ้นไอที' in paragraph.text:
                        paragraph.text = paragraph.text.replace("วันที่เสร็จสิ้นไอที", "{{date_success}}")
                    if 'วันที่ส่งมอบไอที' in paragraph.text:
                        paragraph.text = paragraph.text.replace("วันที่ส่งมอบไอที", "{{date_success_end}}")
                    if 'ลายเซ็นผู้ดำเนินการไอที' in paragraph.text:
                        paragraph.text = paragraph.text.replace("ลายเซ็นผู้ดำเนินการไอที", "{{signature}}")
                    if 'วันรับเรื่อง' in paragraph.text:
                        paragraph.text = paragraph.text.replace("วันรับเรื่อง", "{{date_update}}")
                    if 'ชื่อผู้ทำ' in paragraph.text:
                        paragraph.text = paragraph.text.replace("ชื่อผู้ทำ", "{{name}}")
                    if 'วันที่เซ็น' in paragraph.text:
                        paragraph.text = paragraph.text.replace("วันที่เซ็น", "{{date_sent}}")
                    if 'รวมวันทำงาน' in paragraph.text:
                        paragraph.text = paragraph.text.replace("รวมวันทำงาน", "{{all_day}}")
                        document.save('./static/document/it.docx')
